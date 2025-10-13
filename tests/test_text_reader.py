"""Tests for text_reader module"""
import pytest
from pathlib import Path
import tempfile
import shutil
from docx import Document
from src.text_reader import TextReader


@pytest.fixture
def temp_dir():
    """Create temporary directory"""
    temp_path = Path(tempfile.mkdtemp())
    yield temp_path
    shutil.rmtree(temp_path)


@pytest.fixture
def reader():
    """Create TextReader instance"""
    return TextReader()


class TestTextReader:
    """Tests for TextReader class"""

    def test_initialization(self, reader):
        """Test initialization"""
        assert reader is not None

    def test_read_text_file(self, reader, temp_dir):
        """Test reading text file"""
        test_file = temp_dir / "test.txt"
        content = "Test content\nSecond line"
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(content)

        result = reader.read_text(str(test_file))
        assert result == content

    def test_read_docx_file(self, reader, temp_dir):
        """Test reading docx file"""
        test_file = temp_dir / "test.docx"

        # Create docx with content
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        doc.save(test_file)

        result = reader.read_docx(str(test_file))

        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Third paragraph" in result

    def test_read_markdown_file(self, reader, temp_dir):
        """Test reading markdown file"""
        test_file = temp_dir / "test.md"
        content = "# Heading\n\nParagraph with **bold** and *italic*\n\n- List item 1\n- List item 2"
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write(content)

        result = reader.read_markdown(str(test_file))

        assert "Heading" in result
        assert "Paragraph" in result
        assert "List item" in result

    def test_read_file_auto_detect(self, reader, temp_dir):
        """Test automatic file format detection"""
        # Test .txt
        txt_file = temp_dir / "test.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write("Text content")

        result = reader.read_file(str(txt_file))
        assert "Text content" in result

    def test_read_file_unsupported_format(self, reader, temp_dir):
        """Test error on unsupported format"""
        test_file = temp_dir / "test.xyz"
        test_file.touch()

        with pytest.raises(ValueError, match="Unsupported file format"):
            reader.read_file(str(test_file))

    def test_read_file_not_found(self, reader):
        """Test error when file doesn't exist"""
        with pytest.raises(FileNotFoundError):
            reader.read_file("/nonexistent/file.txt")

    def test_strip_markdown(self, reader):
        """Test markdown stripping"""
        markdown = "# Heading\n\n**Bold** and *italic* text\n\n[Link](url)"
        result = reader._strip_markdown(markdown)

        assert "Heading" in result
        assert "Bold" in result
        assert "italic" in result
        assert "**" not in result
        assert "*" not in result or result.count("*") < markdown.count("*")

    def test_detect_language_russian(self, reader):
        """Test detecting Russian language"""
        text = "Это тестовый текст на русском языке"
        result = reader.detect_language(text)
        assert result == 'ru'

    def test_detect_language_english(self, reader):
        """Test detecting English language"""
        text = "This is a test text in English"
        result = reader.detect_language(text)
        assert result == 'en'
