"""
Module responsible for producing DOCX and Markdown output documents.
"""
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .config import settings
from .logger import logger
from .utils import sanitize_filename


class DocumentWriter:
    """Helper for rendering transcripts and translations into documents."""

    def __init__(self):
        self.output_dir = settings.OUTPUT_DIR

    def create_documents(
        self,
        title: str,
        sections: List[dict],
        original_text: Optional[str] = None,
    ) -> tuple[Path, Path]:
        """
        Produce both DOCX and Markdown documents.

        Args:
            title: Base title for output files.
            sections: List of sections with keys `title`, `method`, and `content`.
            original_text: Optional raw text (currently unused).

        Returns:
            Tuple containing paths to the DOCX and Markdown files.
        """
        clean_title = sanitize_filename(title)

        docx_path = self.output_dir / f"{clean_title}.docx"
        md_path = self.output_dir / f"{clean_title}.md"

        logger.info("Creating documents for title: %s", clean_title)

        self._create_docx(docx_path, title, sections)
        self._create_markdown(md_path, title, sections)

        logger.info("Documents saved:\n  - %s\n  - %s", docx_path, md_path)

        return docx_path, md_path

    def _create_docx(self, path: Path, title: str, sections: List[dict]) -> None:
        """Render a DOCX document."""
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = settings.DEFAULT_FONT
        font.size = Pt(settings.DEFAULT_FONT_SIZE)

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for section in sections:
            doc.add_heading(section["title"], level=1)

            if "method" in section and section["method"]:
                method_para = doc.add_paragraph()
                method_run = method_para.add_run(f"Method: {section['method']}")
                method_run.italic = True
                method_run.font.size = Pt(10)
                method_run.font.color.rgb = RGBColor(128, 128, 128)

            content = section.get("content", "")
            if content:
                paragraphs = content.split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.style = "Normal"

                        # Apply justify alignment for translation and transcription sections
                        section_title = section.get("title", "")
                        if section_title in ["Перевод", "Транскрипция", "Summary"]:
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            doc.add_paragraph()

        doc.save(str(path))

    def _create_markdown(self, path: Path, title: str, sections: List[dict]) -> None:
        """Render a Markdown document."""
        lines: List[str] = []

        lines.append(f"# {title}\n")

        for section in sections:
            lines.append(f"## {section['title']}\n")

            if "method" in section and section["method"]:
                lines.append(f"*Method: {section['method']}*\n")

            content = section.get("content", "")
            if content:
                lines.append(content)
                lines.append("")

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def create_from_segments(
        self,
        title: str,
        transcription_segments: List,
        translation_segments: Optional[List] = None,
        transcribe_method: str = "",
        translate_method: str = "",
        with_timestamps: bool = True,
        with_speakers: bool = False,
        description: str = "",
    ) -> tuple[Path, Path]:
        """
        Build documents directly from transcription/translation segments.

        Args:
            title: Output title.
            transcription_segments: Source transcription segments.
            translation_segments: Optional translation segments.
            transcribe_method: Transcription backend label.
            translate_method: Translation backend label.
            with_timestamps: Whether to include timestamps.
            with_speakers: Whether to include speaker labels.
            description: Document content description (optional).

        Returns:
            Tuple with DOCX and Markdown paths.
        """
        from .transcriber import Transcriber

        transcriber = Transcriber()
        sections: List[dict] = []

        # Add metadata section if description is provided
        if description or transcribe_method or translate_method:
            metadata_lines = []

            if description:
                metadata_lines.append(f"**Содержимое:** {description}")

            if transcribe_method:
                metadata_lines.append(f"**Метод транскрипции:** {transcribe_method}")

            if translation_segments and translate_method:
                metadata_lines.append(f"**Метод перевода:** {translate_method}")

            sections.append({
                "title": "Информация о документе",
                "method": "",
                "content": "\n\n".join(metadata_lines)
            })

        if translation_segments:
            if with_timestamps:
                translation_text = transcriber.segments_to_text_with_timestamps(
                    translation_segments,
                    with_speakers=with_speakers
                )
            else:
                translation_text = transcriber.segments_to_text(translation_segments)

            sections.append(
                {
                    "title": "Перевод",
                    "method": "",
                    "content": translation_text,
                }
            )

        if with_timestamps:
            transcription_text = transcriber.segments_to_text_with_timestamps(
                transcription_segments,
                with_speakers=with_speakers
            )
        else:
            # Even without timestamps, we can show speaker labels
            if with_speakers:
                transcription_text = transcriber.segments_to_text_with_speakers(transcription_segments)
            else:
                transcription_text = transcriber.segments_to_text(transcription_segments)

        sections.append(
            {
                "title": "Транскрипция",
                "method": "",
                "content": transcription_text,
            }
        )

        return self.create_documents(title, sections)
